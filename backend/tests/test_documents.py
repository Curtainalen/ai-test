import io
import pytest
from docx import Document
from app.errors import AppError
from app.services.documents import ai_module_candidates, parse_document, sha256_bytes, suggest_modules, validate_filename

def test_txt_markdown_and_docx_source_locations():
    txt=parse_document("a.txt","第一行\n\n第二行".encode()); md=parse_document("a.md",b"# Login\n- success\n```json\n{}\n```\n| A | B |\n|---|---|\n| 1 | 2 |")
    doc=Document(); doc.add_heading("用户登录",level=1); doc.add_paragraph("输入账号密码"); stream=io.BytesIO(); doc.save(stream); dx=parse_document("a.docx",stream.getvalue())
    assert txt[1]["source_locator"]["line_start"]==3
    assert [b["block_type"] for b in md]==["heading","list","code","table"]
    assert dx[0]["block_type"]=="heading" and dx[0]["source_locator"]["paragraph_index"]==0

def test_filename_hash_and_unsupported_doc():
    assert sha256_bytes(b"x")==sha256_bytes(b"x")
    with pytest.raises(AppError) as caught: validate_filename("legacy.doc")
    assert caught.value.code=="FILE_UNSUPPORTED"
    with pytest.raises(AppError): validate_filename("../a.txt")

def test_heading_and_rule_module_splitting():
    headings = parse_document("a.md", b"# Login\naccount access\n# Orders\ncreate order")
    result = suggest_modules(headings)
    assert [item["name"] for item in result] == ["Login", "Orders"]
    assert all(item["split_method"] == "heading" for item in result)
    rules = parse_document("a.txt", "1. Login\nlogin detail\n2. Orders\norder detail".encode())
    assert len(suggest_modules(rules)) == 2

def test_ai_module_json_validation_rejects_invalid_sources_for_fallback():
    blocks = parse_document("a.txt", b"login")
    with pytest.raises(ValueError):
        ai_module_candidates({"modules": [{"name": "Login", "source_block_sequences": [99]}]}, blocks)
