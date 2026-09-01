from __future__ import annotations
import hashlib, io, re
from pathlib import Path
from docx import Document
from pypdf import PdfReader
from app.errors import AppError

ALLOWED = {".pdf":"application/pdf", ".docx":"application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".md":"text/markdown", ".markdown":"text/markdown", ".txt":"text/plain"}

def validate_filename(filename: str) -> tuple[str,str]:
    safe = Path(filename or "").name
    if safe != filename or not safe: raise AppError("FILE_INVALID_NAME", "文件名无效", 422)
    ext = Path(safe).suffix.lower()
    if ext == ".doc": raise AppError("FILE_UNSUPPORTED", "DOC 暂不支持稳定解析，请转换为 DOCX 或 PDF", 415)
    if ext not in ALLOWED: raise AppError("FILE_UNSUPPORTED", "仅支持 PDF、DOCX、Markdown 和 TXT", 415)
    return safe, ext

def sha256_bytes(content: bytes) -> str: return hashlib.sha256(content).hexdigest()

def parse_document(filename: str, content: bytes, max_pdf_pages: int = 200, max_docx_images: int = 200) -> list[dict]:
    _, ext = validate_filename(filename)
    if ext == ".pdf": return _parse_pdf(content, max_pdf_pages)
    if ext == ".docx": return _parse_docx(content, max_docx_images)
    text = decode_text(content)
    return _parse_markdown(text) if ext in {".md", ".markdown"} else _parse_txt(text)

def decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try: return content.decode(encoding)
        except UnicodeDecodeError: continue
    raise AppError("DOCUMENT_ENCODING_UNSUPPORTED", "文本编码无法识别", 422)

def _block(seq, kind, content, locator, structured=None, confidence=1.0, needs=False):
    return {"seq":seq,"block_type":kind,"content":content,"structured_content":structured or {},"source_locator":locator,"confidence":confidence,"needs_correction":needs}

def _parse_txt(text: str) -> list[dict]:
    blocks=[]
    for i,line in enumerate(text.splitlines(),1):
        value=line.strip()
        if value: blocks.append(_block(len(blocks)+1,"paragraph",value,{"line_start":i,"line_end":i}))
    return blocks

def _parse_markdown(text: str) -> list[dict]:
    blocks=[]; headings=[]; code=[]; in_code=False; lines=text.splitlines(); i=1
    while i <= len(lines):
        line=lines[i-1]
        if line.strip().startswith("```"):
            if in_code: blocks.append(_block(len(blocks)+1,"code","\n".join(code),{"line_start":i-len(code),"line_end":i,"heading_path":headings.copy()})); code=[]
            in_code=not in_code; i+=1; continue
        if in_code: code.append(line); i+=1; continue
        match=re.match(r"^(#{1,6})\s+(.+)$",line)
        if match:
            level=len(match.group(1)); headings=headings[:level-1]+[match.group(2).strip()]
            blocks.append(_block(len(blocks)+1,"heading",headings[-1],{"line_start":i,"line_end":i,"heading_path":headings.copy()},{"level":level})); i+=1; continue
        value=line.strip()
        if not value: i+=1; continue
        if "|" in value and i < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[i]):
            table_lines=[value,lines[i].strip()]; end=i+1
            while end < len(lines) and "|" in lines[end]: table_lines.append(lines[end].strip()); end+=1
            rows=[[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines if not re.match(r"^\s*\|?\s*:?-+",row)]
            blocks.append(_block(len(blocks)+1,"table","\n".join(table_lines),{"line_start":i,"line_end":end,"heading_path":headings.copy()},{"rows":rows})); i=end+1; continue
        kind="list" if re.match(r"^([-*+] |\d+\. )",value) else "paragraph"
        blocks.append(_block(len(blocks)+1,kind,value,{"line_start":i,"line_end":i,"heading_path":headings.copy()}))
        i+=1
    return blocks

def _parse_pdf(content: bytes, max_pages: int) -> list[dict]:
    try: reader=PdfReader(io.BytesIO(content))
    except Exception as exc: raise AppError("DOCUMENT_PARSE_FAILED", "PDF 无法解析", 422) from exc
    if reader.is_encrypted: raise AppError("DOCUMENT_PASSWORD_REQUIRED", "加密 PDF 暂不接受解析密码，请上传解密版本", 422)
    if len(reader.pages)>max_pages: raise AppError("DOCUMENT_PAGE_LIMIT", "PDF 页数超过限制", 413, {"pages":len(reader.pages),"limit":max_pages})
    blocks=[]
    for page_no,page in enumerate(reader.pages,1):
        text=(page.extract_text() or "").strip(); confidence=1.0 if text else 0.0
        if text: blocks.append(_block(len(blocks)+1,"paragraph",text,{"page":page_no},confidence=confidence))
        else: blocks.append(_block(len(blocks)+1,"image","",{"page":page_no,"ocr_status":"not_implemented"},confidence=0.0,needs=True))
    return blocks

def _parse_docx(content: bytes, max_images: int) -> list[dict]:
    try: doc=Document(io.BytesIO(content))
    except Exception as exc: raise AppError("DOCUMENT_PARSE_FAILED", "DOCX 无法解析", 422) from exc
    if len(doc.inline_shapes)>max_images: raise AppError("DOCUMENT_IMAGE_LIMIT", "DOCX 图片数量超过限制", 413)
    blocks=[]
    for index,p in enumerate(doc.paragraphs):
        text=p.text.strip()
        if not text: continue
        style=(p.style.name or "").lower(); kind="heading" if style.startswith("heading") else ("list" if "list" in style else "paragraph")
        blocks.append(_block(len(blocks)+1,kind,text,{"paragraph_index":index}, {"style":p.style.name}))
    for index,table in enumerate(doc.tables):
        rows=[[cell.text for cell in row.cells] for row in table.rows]
        blocks.append(_block(len(blocks)+1,"table","\n".join(" | ".join(row) for row in rows),{"table_index":index},{"rows":rows}))
    for index,_shape in enumerate(doc.inline_shapes):
        blocks.append(_block(len(blocks)+1,"image","",{"image_index":index,"ocr_status":"not_implemented"},confidence=0.0,needs=True))
    return blocks

def suggest_modules(blocks: list[dict]) -> list[dict]:
    headings=[(index, block) for index, block in enumerate(blocks) if block["block_type"]=="heading"]
    if headings:
        result=[]
        for index, heading in headings:
            level=(heading.get("structured_content") or {}).get("level", 6)
            end=len(blocks)
            for next_index in range(index + 1, len(blocks)):
                candidate=blocks[next_index]
                if candidate["block_type"] == "heading" and (candidate.get("structured_content") or {}).get("level", 6) <= level:
                    end=next_index; break
            selected=blocks[index:end]
            paragraphs=[block["content"].strip() for block in selected[1:] if block["block_type"] in {"paragraph", "list"} and block["content"].strip()]
            result.append({"name":heading["content"][:255], "description":" ".join(paragraphs)[:1000], "source_block_ids":[], "source_seqs":[block["seq"] for block in selected], "split_method":"heading"})
        return result
    if not blocks: return []
    boundaries=[index for index, block in enumerate(blocks) if re.match(r"^(?:\d+(?:\.\d+)*|REQ[-_A-Z0-9]+)[、.\s-]+", block["content"].strip(), re.I)]
    if boundaries:
        result=[]
        for offset, start in enumerate(boundaries):
            selected=blocks[start:boundaries[offset + 1] if offset + 1 < len(boundaries) else len(blocks)]
            result.append({"name":selected[0]["content"][:255], "description":" ".join(block["content"] for block in selected[1:] if block["content"])[:1000], "source_block_ids":[], "source_seqs":[block["seq"] for block in selected], "split_method":"rule"})
        return result
    return [{"name":"需求模块 1","description":"","source_block_ids":[],"source_seqs":[block["seq"] for block in blocks], "split_method":"rule"}]


def ai_module_candidates(payload: object, blocks: list[dict]) -> list[dict]:
    """Validate structured AI output against only the supplied document blocks."""
    if not isinstance(payload, dict) or not isinstance(payload.get("modules"), list):
        raise ValueError("AI module payload must contain modules")
    known = {block["seq"] for block in blocks}
    result=[]
    for item in payload["modules"]:
        if not isinstance(item, dict) or not isinstance(item.get("name"), str) or not item["name"].strip():
            raise ValueError("AI module name is invalid")
        seqs=item.get("source_block_sequences")
        if not isinstance(seqs, list) or not seqs or any(not isinstance(seq, int) or seq not in known for seq in seqs):
            raise ValueError("AI module references invalid source blocks")
        confidence=item.get("confidence")
        if confidence is not None and (not isinstance(confidence, (int, float)) or not 0 <= confidence <= 1):
            raise ValueError("AI confidence is invalid")
        result.append({"name":item["name"].strip()[:255], "description":str(item.get("description") or "")[:1000], "source_seqs":list(dict.fromkeys(seqs)), "split_method":"ai", "confidence":confidence})
    return result
